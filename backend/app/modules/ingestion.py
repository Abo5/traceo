"""Requirements Parser (TRD §4.1, FR-REQ).

Pipeline: upload -> async job -> text extraction (pdf/docx/md/txt) -> deterministic
segmentation -> per-segment `extract_requirement` LLM call -> persist Requirements
with provenance.

Re-upload of an existing filename bumps the document version and diffs the extraction
against the previous inventory by external_id, then content_hash (FR-REQ-06, FR-TRC-04).
"""
import hashlib
import re
import uuid
from pathlib import Path

from fastapi import APIRouter, Body, Depends, File, HTTPException, UploadFile
from sqlalchemy import func, or_
from sqlalchemy.orm import Session

from .. import jobs as jobstore
from ..config import settings
from ..db import SessionLocal, get_db
from ..deps import audit, get_project_scoped, require
from ..llm import UNTRUSTED_NOTE, frame_untrusted, get_provider
from ..models import Project, Requirement, RequirementTestCase, SourceDocument, User
from .generation import try_autopilot_generation

router = APIRouter()

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
MAX_SEGMENTS = 500
MIN_SEGMENT_CHARS = 15

REQUIREMENT_TYPES = {"functional", "business_rule", "data", "interface", "non_functional"}

# The segment comes from a file the user uploaded, so it is untrusted input: it is
# framed by llm.frame_untrusted and introduced by UNTRUSTED_NOTE before the
# "SEGMENT:\n" sentinel. The sentinel itself is unchanged — MockProvider splits on
# it (app/llm/mock.py) and strips the frame, so the offline path is unaffected.
EXTRACT_PROMPT = (
    "Extract the software requirement from this segment. "
    "Answer in English.\n"
    + UNTRUSTED_NOTE
    + "SEGMENT:\n"
)

EXTRACT_SCHEMA = {
    "type": "object",
    "properties": {
        "external_id": {"type": "string"},
        "description": {"type": "string"},
        "acceptance_criteria": {"type": "array", "items": {"type": "string"}},
        "type": {"enum": sorted(REQUIREMENT_TYPES)},
        "priority": {"type": "string"},
        "confidence": {"type": "number"},
    },
    "required": ["external_id", "description", "acceptance_criteria",
                 "type", "priority", "confidence"],
}

# --- deterministic text utilities -------------------------------------------------

# Requirement-ID openers: REQ-1 / FR-01 / BR_2 / NFR 3 / numbered clauses "3.1.2"
REQ_ID_LINE = re.compile(
    r"^\s*(?:"
    r"(?:REQ|FR|BR|NFR|UC|SRS|BUS)[-_ ]?\d+(?:[.-]\d+)*"
    r"|\d+(?:\.\d+)+"
    r")\b[.:)\-–—]?",
    re.IGNORECASE,
)
HEADING_RE = re.compile(r"^\s*#{1,6}\s+\S")
BULLET_RE = re.compile(r"^\s*(?:[-*•▪◦]|\d+[.)]|[a-h][.)])\s+\S")


def _content_hash(description: str, criteria: list) -> str:
    payload = (description or "") + "\n" + "\n".join(str(c) for c in (criteria or []))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


# --- text extraction ---------------------------------------------------------------

def _extract_text(path: Path, ext: str) -> list[tuple[int | None, str]]:
    """Return list of (page_number, text). page_number is None when unknowable."""
    if ext == ".pdf":
        try:
            import fitz  # pymupdf
        except ImportError:
            raise RuntimeError(
                "PDF support requires the 'pymupdf' package (pip install pymupdf); "
                "it is not installed on this server."
            )
        pages: list[tuple[int | None, str]] = []
        with fitz.open(path) as doc:
            for page_index, page in enumerate(doc):
                pages.append((page_index + 1, page.get_text("text") or ""))
        return pages

    if ext == ".docx":
        try:
            import docx  # python-docx
        except ImportError:
            raise RuntimeError(
                "DOCX support requires the 'python-docx' package (pip install python-docx); "
                "it is not installed on this server."
            )
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return [(None, "\n".join(parts))]

    if ext in (".md", ".txt"):
        return [(None, path.read_text(encoding="utf-8", errors="replace"))]

    raise RuntimeError(f"Unsupported file extension '{ext}'")


# --- deterministic segmentation ----------------------------------------------------

def segment_pages(pages: list[tuple[int | None, str]]) -> list[dict]:
    """Split extracted text into candidate requirement segments.

    Boundaries: requirement-ID lines, markdown headings, blank-line paragraph breaks.
    A requirement line keeps its following bullet lines (acceptance criteria) attached,
    even across single blank lines. Segments < MIN_SEGMENT_CHARS are skipped; capped
    at MAX_SEGMENTS. Fully deterministic — runs before any LLM call.
    """
    segments: list[dict] = []
    cur_lines: list[str] = []
    cur_page: int | None = None

    def flush():
        nonlocal cur_lines
        text = "\n".join(cur_lines).strip()
        if len(text) >= MIN_SEGMENT_CHARS:
            segments.append({"text": text, "page": cur_page})
        cur_lines = []

    pending_blank = False
    for page_no, text in pages:
        for raw in text.splitlines():
            line = raw.rstrip()
            if not line.strip():
                pending_blank = True
                continue
            is_req = bool(REQ_ID_LINE.match(line))
            is_heading = bool(HEADING_RE.match(line))
            is_bullet = (not is_req) and bool(BULLET_RE.match(line))
            if is_req or is_heading:
                flush()
                cur_page = page_no
                cur_lines = [line]
            elif is_bullet and cur_lines:
                # acceptance-criteria bullets stay grouped with their requirement line
                cur_lines.append(line)
            elif pending_blank and cur_lines:
                flush()
                cur_page = page_no
                cur_lines = [line]
            else:
                if not cur_lines:
                    cur_page = page_no
                cur_lines.append(line)
            pending_blank = False
    flush()
    return segments[:MAX_SEGMENTS]


# --- LLM structuring (per-segment, failure-isolated) --------------------------------

def _structure_segment(provider, segment_text: str) -> dict:
    """One LLM call per segment. A failing segment degrades to a raw-text requirement
    with confidence 0.3 — it is never silently dropped."""
    try:
        result = provider.complete_json(
            "extract_requirement", EXTRACT_PROMPT + frame_untrusted(segment_text),
            EXTRACT_SCHEMA)
        data = dict(result.data)
    except Exception:
        return {
            "external_id": "",
            "description": segment_text[:2000],
            "acceptance_criteria": [],
            "type": "functional",
            "priority": "medium",
            "confidence": 0.3,
        }
    data["external_id"] = str(data.get("external_id") or "").strip()
    data["description"] = str(data.get("description") or "").strip() or segment_text[:2000]
    data["acceptance_criteria"] = [str(c) for c in (data.get("acceptance_criteria") or [])]
    if data.get("type") not in REQUIREMENT_TYPES:
        data["type"] = "functional"
    data["priority"] = str(data.get("priority") or "medium")
    try:
        data["confidence"] = max(0.0, min(1.0, float(data.get("confidence", 0.5))))
    except (TypeError, ValueError):
        data["confidence"] = 0.5
    return data


def confirm_all_extracted(db: Session, org_id: str, project_id: str) -> int:
    """Flip every 'extracted' requirement to 'confirmed'; returns how many.

    The single code path shared by the manual confirm_all endpoint and the
    autopilot chain (contract 4a) — auditing is the caller's responsibility
    because the action name differs (requirement.confirm_all vs
    auto.requirements.confirm_all)."""
    reqs = db.query(Requirement).filter(
        Requirement.project_id == project_id,
        Requirement.organisation_id == org_id,
        Requirement.state == "extracted").all()
    for r in reqs:
        r.state = "confirmed"
    return len(reqs)


def _mark_stale(db: Session, requirement_id: str) -> None:
    """Lazy import to avoid a circular dependency with the traceability module."""
    try:
        from .traceability import mark_stale
    except ImportError:
        return
    mark_stale(db, requirement_id)


# --- persistence with re-upload diffing ---------------------------------------------

def _persist_requirements(db: Session, doc: SourceDocument, extractions: list[dict]) -> dict:
    """Insert/diff extracted requirements.

    First upload: everything inserted state='extracted'.
    Re-upload (same filename): match to existing project requirements by external_id
    first, then by content_hash. Unchanged -> keep; changed -> update in place,
    version += 1, state='changed', mark linked approved cases stale; existing rows from
    the previous document version left unmatched -> state='removed'; new -> insert.
    """
    counts = {"added": 0, "changed": 0, "unchanged": 0, "removed": 0}

    prior_doc_ids = [
        d.id for d in db.query(SourceDocument)
        .filter(SourceDocument.project_id == doc.project_id,
                SourceDocument.organisation_id == doc.organisation_id,
                SourceDocument.filename == doc.filename,
                SourceDocument.id != doc.id).all()
    ]
    existing: list[Requirement] = []
    if prior_doc_ids:
        existing = db.query(Requirement).filter(
            Requirement.project_id == doc.project_id,
            Requirement.organisation_id == doc.organisation_id,
            Requirement.source_document_id.in_(prior_doc_ids),
            Requirement.state != "removed",
        ).all()

    by_external_id = {r.external_id: r for r in existing if r.external_id}
    by_hash = {r.content_hash: r for r in existing if r.content_hash}
    matched_ids: set[str] = set()

    for item in extractions:
        data = item["data"]
        new_hash = _content_hash(data["description"], data["acceptance_criteria"])
        location = {"page": item["page"], "index": item["index"]}

        match = None
        if data["external_id"] and data["external_id"] in by_external_id:
            candidate = by_external_id[data["external_id"]]
            if candidate.id not in matched_ids:
                match = candidate
        if match is None and new_hash in by_hash and by_hash[new_hash].id not in matched_ids:
            match = by_hash[new_hash]

        if match is not None:
            matched_ids.add(match.id)
            match.source_document_id = doc.id
            match.source_text = item["segment"]
            match.source_location = location
            if match.content_hash == new_hash:
                counts["unchanged"] += 1
            else:
                match.external_id = data["external_id"] or match.external_id
                match.description = data["description"]
                match.acceptance_criteria = data["acceptance_criteria"]
                match.type = data["type"]
                match.priority = data["priority"]
                match.confidence = data["confidence"]
                match.content_hash = new_hash
                match.version += 1
                match.state = "changed"
                _mark_stale(db, match.id)
                counts["changed"] += 1
        else:
            db.add(Requirement(
                organisation_id=doc.organisation_id,
                project_id=doc.project_id,
                source_document_id=doc.id,
                external_id=data["external_id"],
                description=data["description"],
                acceptance_criteria=data["acceptance_criteria"],
                type=data["type"],
                priority=data["priority"],
                state="extracted",
                version=1,
                source_location=location,
                source_text=item["segment"],
                confidence=data["confidence"],
                content_hash=new_hash,
            ))
            counts["added"] += 1

    for req in existing:
        if req.id not in matched_ids:
            req.state = "removed"
            counts["removed"] += 1

    return counts


def _run_ingest(job, document_id: str, project_id: str, org_id: str, actor_id: str):
    """Job body — owns its own session (runs on a worker thread)."""
    db = SessionLocal()
    try:
        doc = db.get(SourceDocument, document_id)
        if doc is None:
            raise RuntimeError("Source document disappeared before parsing")
        doc.parse_status = "parsing"
        db.commit()

        path = settings.STORAGE_DIR / doc.storage_key
        ext = Path(doc.filename).suffix.lower()
        try:
            pages = _extract_text(path, ext)
        except Exception as exc:
            doc.parse_status = "failed"
            doc.parse_error = str(exc)
            db.commit()
            raise

        segments = segment_pages(pages)
        job.message = f"Segmented document into {len(segments)} candidate requirements"

        provider = get_provider()
        extractions: list[dict] = []
        total = len(segments)
        for i, seg in enumerate(segments):
            job.progress = i / total if total else 1.0
            job.message = f"Extracting requirement {i + 1}/{total}"
            data = _structure_segment(provider, seg["text"])
            extractions.append({"data": data, "segment": seg["text"],
                                "page": seg["page"], "index": i})

        job.message = "Persisting requirements"
        counts = _persist_requirements(db, doc, extractions)

        doc.parse_status = "parsed"
        doc.parse_error = None
        audit(db, org_id, actor_id, "document.parsed", "source_document", doc.id,
              {"filename": doc.filename, "version": doc.version,
               "segments": total, **counts})

        result = {"document_id": doc.id, "segments": total, **counts}
        # sessions run with autoflush=False — flush so the freshly persisted
        # requirements are visible to the autopilot chain's queries below
        db.flush()
        project = db.get(Project, project_id)

        # -- autopilot chain (contract 4a): confirm ALL extracted requirements,
        #    then try the generation trigger (4b). Auto stops at draft cases —
        #    approval and runs stay manual (BO-07).
        automation_on = project is not None and project.automation == "auto"
        if automation_on:
            job.message = "Autopilot: confirming extracted requirements"
            confirmed = confirm_all_extracted(db, org_id, project_id)
            audit(db, org_id, actor_id, "auto.requirements.confirm_all", "project",
                  project_id, {"count": confirmed})
            result["auto_confirmed"] = confirmed
        db.commit()

        if automation_on:
            gen_job_id = try_autopilot_generation(db, org_id, actor_id, project_id)
            if gen_job_id:
                result["generation_job_id"] = gen_job_id
        return result
    finally:
        db.close()


# --- serializers ---------------------------------------------------------------------

def _doc_dict(d: SourceDocument) -> dict:
    return {
        "id": d.id, "project_id": d.project_id, "filename": d.filename,
        "mime_type": d.mime_type, "size": d.size, "language": d.language,
        "version": d.version, "parse_status": d.parse_status, "parse_error": d.parse_error,
        "created_at": d.created_at.isoformat() if d.created_at else None,
    }


def _req_dict(r: Requirement) -> dict:
    return {
        "id": r.id, "project_id": r.project_id, "source_document_id": r.source_document_id,
        "external_id": r.external_id, "description": r.description,
        "acceptance_criteria": r.acceptance_criteria, "type": r.type,
        "priority": r.priority, "state": r.state, "version": r.version,
        "source_location": r.source_location, "source_text": r.source_text,
        "confidence": r.confidence, "content_hash": r.content_hash,
        "created_at": r.created_at.isoformat() if r.created_at else None,
        "updated_at": r.updated_at.isoformat() if r.updated_at else None,
    }


def _get_requirement(requirement_id: str, user: User, db: Session) -> Requirement:
    req = db.get(Requirement, requirement_id)
    if not req or req.organisation_id != user.organisation_id:
        raise HTTPException(404, detail={"code": "not_found", "message": "Requirement not found"})
    return req


# --- routes ---------------------------------------------------------------------------

@router.post("/projects/{project_id}/documents", status_code=202)
async def upload_document(project_id: str, file: UploadFile = File(...),
                          user: User = Depends(require("upload_documents")),
                          db: Session = Depends(get_db)):
    project = get_project_scoped(project_id, user, db)
    filename = file.filename or "document.txt"
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(422, detail={
            "code": "unsupported_file_type",
            "message": f"Unsupported file type '{ext}'. Allowed: pdf, docx, md, txt."})

    content = await file.read()
    if len(content) > settings.MAX_UPLOAD_MB * 1024 * 1024:
        raise HTTPException(413, detail={
            "code": "file_too_large",
            "message": f"File exceeds the {settings.MAX_UPLOAD_MB}MB upload limit."})
    if not content:
        raise HTTPException(422, detail={"code": "empty_file", "message": "Uploaded file is empty."})

    safename = re.sub(r"[^\w.\-]", "_", filename)[:120]
    storage_key = f"{uuid.uuid4()}_{safename}"
    (settings.STORAGE_DIR / storage_key).write_bytes(content)

    prior_max = db.query(func.max(SourceDocument.version)).filter(
        SourceDocument.project_id == project_id,
        SourceDocument.organisation_id == user.organisation_id,
        SourceDocument.filename == filename,
    ).scalar() or 0

    doc = SourceDocument(
        organisation_id=user.organisation_id, project_id=project_id,
        filename=filename, mime_type=file.content_type or "", size=len(content),
        storage_key=storage_key, language="en",
        version=prior_max + 1, parse_status="pending",
    )
    db.add(doc)
    db.commit()

    doc_id, org_id, actor_id = doc.id, user.organisation_id, user.id
    job = jobstore.submit(
        "ingest", lambda j: _run_ingest(j, doc_id, project_id, org_id, actor_id))
    return {"job_id": job.id, "document_id": doc_id}


@router.get("/projects/{project_id}/documents")
def list_documents(project_id: str, user: User = Depends(require("view")),
                   db: Session = Depends(get_db)):
    get_project_scoped(project_id, user, db)
    docs = db.query(SourceDocument).filter(
        SourceDocument.project_id == project_id,
        SourceDocument.organisation_id == user.organisation_id,
    ).order_by(SourceDocument.created_at.desc()).all()
    return [_doc_dict(d) for d in docs]


@router.get("/projects/{project_id}/requirements")
def list_requirements(project_id: str, state: str = "", type: str = "",
                      priority: str = "", q: str = "",
                      user: User = Depends(require("view")),
                      db: Session = Depends(get_db)):
    get_project_scoped(project_id, user, db)
    query = db.query(Requirement).filter(
        Requirement.project_id == project_id,
        Requirement.organisation_id == user.organisation_id)
    if state:
        query = query.filter(Requirement.state == state)
    if type:
        query = query.filter(Requirement.type == type)
    if priority:
        query = query.filter(Requirement.priority == priority)
    if q:
        needle = f"%{q}%"
        query = query.filter(or_(Requirement.description.ilike(needle),
                                 Requirement.external_id.ilike(needle),
                                 Requirement.source_text.ilike(needle)))
    if state == "extracted":
        # review queue: lowest-confidence extractions surface first (FR-REQ-08)
        query = query.order_by(Requirement.confidence.asc(), Requirement.created_at.asc())
    else:
        query = query.order_by(Requirement.created_at.asc())
    return [_req_dict(r) for r in query.all()]


@router.patch("/requirements/{requirement_id}")
def update_requirement(requirement_id: str, body: dict = Body(...),
                       user: User = Depends(require("edit_requirements")),
                       db: Session = Depends(get_db)):
    req = _get_requirement(requirement_id, user, db)
    changes: dict = {}

    if "state" in body and body["state"] not in (None, "confirmed"):
        raise HTTPException(422, detail={
            "code": "invalid_state",
            "message": "Only state='confirmed' may be set via this endpoint."})
    if "type" in body and body["type"] is not None and body["type"] not in REQUIREMENT_TYPES:
        raise HTTPException(422, detail={
            "code": "invalid_type",
            "message": f"type must be one of {sorted(REQUIREMENT_TYPES)}"})
    if "acceptance_criteria" in body and body["acceptance_criteria"] is not None:
        if not isinstance(body["acceptance_criteria"], list):
            raise HTTPException(422, detail={
                "code": "invalid_criteria", "message": "acceptance_criteria must be a list."})
        body["acceptance_criteria"] = [str(c) for c in body["acceptance_criteria"]]

    content_changed = False
    if body.get("description") is not None and body["description"] != req.description:
        changes["description"] = {"from": req.description, "to": body["description"]}
        req.description = str(body["description"])
        content_changed = True
    if body.get("acceptance_criteria") is not None \
            and body["acceptance_criteria"] != req.acceptance_criteria:
        changes["acceptance_criteria"] = True
        req.acceptance_criteria = body["acceptance_criteria"]
        content_changed = True
    if body.get("external_id") is not None and body["external_id"] != req.external_id:
        changes["external_id"] = {"from": req.external_id, "to": body["external_id"]}
        req.external_id = str(body["external_id"]).strip()
    if body.get("type") is not None and body["type"] != req.type:
        changes["type"] = body["type"]
        req.type = body["type"]
    if body.get("priority") is not None and body["priority"] != req.priority:
        changes["priority"] = body["priority"]
        req.priority = str(body["priority"])

    if content_changed:
        req.content_hash = _content_hash(req.description, req.acceptance_criteria)
        if req.state == "confirmed":
            # editing a confirmed requirement invalidates its approved cases (FR-TRC-04)
            req.version += 1
            _mark_stale(db, req.id)

    if body.get("state") == "confirmed" and req.state != "confirmed":
        changes["state"] = "confirmed"
        req.state = "confirmed"

    if changes:
        audit(db, user.organisation_id, user.id, "requirement.updated",
              "requirement", req.id, {"changes": sorted(changes.keys())})
    db.commit()
    return _req_dict(req)


@router.post("/requirements", status_code=201)
def create_requirement(body: dict = Body(...),
                       user: User = Depends(require("edit_requirements")),
                       db: Session = Depends(get_db)):
    project_id = body.get("project_id") or ""
    description = str(body.get("description") or "").strip()
    if not project_id or not description:
        raise HTTPException(422, detail={
            "code": "missing_fields", "message": "project_id and description are required."})
    get_project_scoped(project_id, user, db)
    rtype = body.get("type") or "functional"
    if rtype not in REQUIREMENT_TYPES:
        raise HTTPException(422, detail={
            "code": "invalid_type", "message": f"type must be one of {sorted(REQUIREMENT_TYPES)}"})
    criteria = [str(c) for c in (body.get("acceptance_criteria") or [])]

    req = Requirement(
        organisation_id=user.organisation_id, project_id=project_id,
        source_document_id=None,
        external_id=str(body.get("external_id") or "").strip(),
        description=description, acceptance_criteria=criteria,
        type=rtype, priority=str(body.get("priority") or "medium"),
        state="confirmed",  # human-authored — no extraction review needed
        version=1, source_location={}, source_text=description,
        confidence=1.0, content_hash=_content_hash(description, criteria),
    )
    db.add(req)
    db.flush()
    audit(db, user.organisation_id, user.id, "requirement.created",
          "requirement", req.id, {"manual": True})
    db.commit()
    return _req_dict(req)


@router.delete("/requirements/{requirement_id}")
def delete_requirement(requirement_id: str,
                       user: User = Depends(require("edit_requirements")),
                       db: Session = Depends(get_db)):
    req = _get_requirement(requirement_id, user, db)
    db.query(RequirementTestCase).filter(
        RequirementTestCase.requirement_id == req.id).delete(synchronize_session=False)
    audit(db, user.organisation_id, user.id, "requirement.deleted",
          "requirement", req.id, {"external_id": req.external_id})
    db.delete(req)
    db.commit()
    return {"deleted": True, "id": requirement_id}


@router.post("/projects/{project_id}/requirements/confirm_all")
def confirm_all_requirements(project_id: str,
                             user: User = Depends(require("edit_requirements")),
                             db: Session = Depends(get_db)):
    get_project_scoped(project_id, user, db)
    count = confirm_all_extracted(db, user.organisation_id, project_id)
    audit(db, user.organisation_id, user.id, "requirement.confirm_all",
          "project", project_id, {"count": count})
    db.commit()
    return {"confirmed": count}
